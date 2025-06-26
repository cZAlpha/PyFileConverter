import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter.font import Font
import ttkbootstrap as ttk
import os  # For file accessing
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx2pdf import convert
import pypandoc  # For DOCX -> TXT
from PIL import Image
from docx import Document
import re
import tempfile
import shutil
from tkinter import ttk as tk_ttk

# Centralized UI color variables
UI_COLORS = {
    'primary': '#f78b6f',  # Bootstrap primary color
    'secondary': '#f25f4c',  # Bootstrap secondary color
    'success': '#28a745',  # Bootstrap success color
    'danger': '#dc3545',  # Bootstrap danger color
    'warning': '#ffc107',  # Bootstrap warning color
    'info': '#17a2b8',  # Bootstrap info color
    'light': '#f8f9fa',  # Bootstrap light color
    'dark': '#343a40'  # Bootstrap dark color
}



class FileConverterApp:
    def __init__(self, root):
        self.root = root
        self.current_theme = "united"  # Default theme
        
        # Create temporary directory for converted files
        self.temp_dir = tempfile.mkdtemp()
        self.converted_files = {}  # Store mapping of original file to converted file path
        
        # Set the application icon
        self.root.iconphoto(False, tk.PhotoImage(file="assets/app_icon/png/pyfileconverter_icon.png"))
        
        # Create list for files
        self.selected_file_list = []
        
        # Establish the valid filetypes
        self.valid_extensions = ['.bmp', '.jpg', '.jpeg', '.png', '.heic', '.pdf', '.txt', '.docx', '.csv']
        
        # Initialize the style
        self.style = ttk.Style()
        
        # Define a font
        self.custom_font = Font(family="Miracode", size=12)  # Customize as needed
        self.custom_font_title = Font(family="Miracode", size=16)  # Customize as needed
        
        # Define styles for buttons and radiobuttons
        self.style.configure('Custom.TButton', font=self.custom_font, background=UI_COLORS['primary'],
                            foreground=UI_COLORS['light'])
        self.style.map('Custom.TButton', background=[('active', UI_COLORS['secondary'])])
        
        self.style.configure('Custom.TRadiobutton', font=self.custom_font, background=UI_COLORS['light'],
                            foreground=UI_COLORS['light'])
        self.style.map('Custom.TRadiobutton', background=[('active', UI_COLORS['primary'])])
        
        # Define a custom style for settings button with larger font size
        self.style.configure('Large.TButton', font=('Miracode', 18), background=UI_COLORS['primary'],
                            foreground=UI_COLORS['light'])
        self.style.map('Large.TButton', background=[('active', UI_COLORS['secondary'])])
        
        # Create main and settings frames
        self.main_frame = ttk.Frame(root)
        self.settings_frame = ttk.Frame(root)
        
        # Pack main frame and initialize settings frame
        self.main_frame.pack(fill=tk.BOTH, expand=True)
        self.settings_frame.pack(fill=tk.BOTH, expand=True)
        
        # Create widgets for the main frame
        self.create_main_widgets()
        
        # Create widgets for the settings frame
        self.create_settings_widgets()
        
        # Set initial theme
        self.set_theme(self.current_theme)
        
        # Dictionary to store delete buttons with their item IDs
        self.delete_buttons = {}
        
        # Bind the resizing of the window and allat
        self.root.bind('<Configure>', self.on_configure)
    
    def create_main_widgets(self):
        # File conversion type options
        self.conversion_type_var = tk.StringVar(value="Select A Filetype")  # Default type
        
        # Create the title label and place it at the top center
        ttk.Label(self.main_frame, text="PyFile Converter", font=self.custom_font_title).grid(row=0, column=1, pady=10,
                                                                                            padx=(0, 60), sticky="n")
        # Create a frame to house the clear all button
        top_left_frame = ttk.Frame(self.main_frame)
        top_left_frame.grid(row=0, column=0, padx=10, pady=10, sticky="nw")
        
        # Place the clear button
        clear_button = ttk.Button(top_left_frame, text="Clear All", command=self.clear_all_entries, style='danger', width=8)
        clear_button.pack(side=tk.RIGHT)
        
        # Create a frame to hold the settings button in the top-right corner
        top_right_frame = ttk.Frame(self.main_frame)
        top_right_frame.grid(row=0, column=2, padx=10, pady=10, sticky="ne")
        
        # Place the Settings button in the top-right frame using the custom style with larger font
        settings_button = ttk.Button(top_right_frame, text="⚙", command=self.show_settings, style='secondary')
        settings_button.pack(side=tk.RIGHT)
        
        # Create a frame to hold the import and convert buttons in the middle-rightmost sector
        middle_right_frame = ttk.Frame(self.main_frame)
        middle_right_frame.grid(row=1, column=2, padx=(0, 40), pady=10, sticky="ne")
        
        # Create the Import button
        import_button = ttk.Button(middle_right_frame, text="Import Files", command=self.on_file_import,
                                    style='Custom.TButton')
        import_button.pack(side=tk.TOP, pady=10)
        
        # Create the Convert button
        convert_button = ttk.Button(middle_right_frame, text="Convert Files", command=self.convert_all_files,
                                    style='Custom.TButton')
        convert_button.pack(side=tk.TOP, pady=5)
        
        # Create the Download All button
        download_all_button = ttk.Button(middle_right_frame, text="Download All", command=self.download_all_files,
                                        style='success.TButton')
        download_all_button.pack(side=tk.BOTTOM, pady=5)
        
        # Create a frame for file listing
        self.file_list_frame = ttk.Frame(self.main_frame)
        self.file_list_frame.grid(row=1, column=0, columnspan=2, pady=10, padx=10, sticky="nsew")
        
        # Create Treeview for file list
        self.file_list = ttk.Treeview(self.file_list_frame, columns=("No", "File Name", "Additional Info", 
                                                                    "Download", "Delete"),
                                                            show="headings")
        self.file_list.heading("#1", text="#", anchor="e")  # Number of files heading
        self.file_list.heading("#2", text="File Name", anchor="w")  # File name heading
        self.file_list.heading("#3", text="Output Filetype", anchor="w")  # Additional information heading
        self.file_list.heading("#4", text="Download", anchor="w")  # Download column heading
        self.file_list.heading("#5", text="Delete", anchor="w")  # Delete column heading
        self.file_list.column("#1", width=30, minwidth=20, anchor="e", stretch=False)  # Number of files column
        self.file_list.column("#2", width=200, anchor="w")  # File name column
        self.file_list.column("#3", width=150, anchor="w")  # Additional information column
        self.file_list.column("#4", width=100, anchor="w")  # Download column
        self.file_list.column("#5", width=60, anchor="w")  # Delete column
        
        # Pack the Treeview with padding on the right side
        self.file_list.pack(expand=True, fill=tk.BOTH, padx=20, pady=20)  # Add 20px padding on the right side
        
        # Dictionary to store dropdown widgets and download buttons
        self.dropdown_widgets = {}
        self.download_buttons = {}
        
        # Configure grid column weights to ensure proper alignment
        self.main_frame.grid_columnconfigure(0, weight=0, minsize=10)  # Skinny leftmost column
        self.main_frame.grid_columnconfigure(2, weight=1)
        self.main_frame.grid_rowconfigure(2, weight=1)  # Ensure the file list expands
        
        # Bind the Treeview select event to start editing
        self.file_list.bind("<Double-1>", self.on_item_double_click)
    
    def on_item_double_click(self, event):
        item = self.file_list.selection()
        if not item:
            return
        
        column = self.file_list.identify_column(event.x)
        column = column.split('#')[1]
        
        # Adjust to handle the correct column, here we assume column 4 is "Additional Info"
        if int(column) == 3:
            self.start_editing(item[0], int(column))
    
    def start_editing(self, item_id, column):
        # Remove any existing Entry widget
        if hasattr(self, 'entry'):
            self.entry.destroy()
        
        # Get the coordinates where the click occurred
        x, y, width, height = self.file_list.bbox(item_id, column)
        
        # Create and place an Entry widget
        self.entry = tk.Entry(self.file_list_frame)
        self.entry.place(x=x + self.file_list_frame.winfo_rootx(), y=y + self.file_list_frame.winfo_rooty(),
                        width=width)
        
        # Set the current value in the Entry
        self.entry.insert(0, self.file_list.item(item_id, 'values')[column - 1])
        
        # Focus on the Entry widget and bind the return key to save the edit
        self.entry.focus_set()
        self.entry.bind('<Return>', lambda e: self.save_edit(item_id, column))
    
    def save_edit(self, item_id, column):
        # Save the new value and update the Treeview
        new_value = self.entry.get()
        self.file_list.item(item_id, values=self.update_item_values(item_id, column, new_value))
        self.entry.destroy()
    
    def update_item_values(self, item_id, column, new_value):
        values = list(self.file_list.item(item_id, 'values'))
        values[column - 1] = new_value
        return values
    
    def create_settings_widgets(self):
        # Widgets for settings
        ttk.Label(self.settings_frame, text="Select Theme:", font=self.custom_font_title).pack(pady=10)
        
        self.theme_var = tk.StringVar(value=self.current_theme)
        
        # Dark  mode button
        darkmodeButton = ttk.Button(self.settings_frame, text="Dark Mode", command=self.set_theme_dark, style='Custom.LargeTButton')
        darkmodeButton.pack(pady=5)
        # Light mode button
        lightmodeButton = ttk.Button(self.settings_frame, text="Light Mode", command=self.set_theme_light, style='Custom.LargeTButton')
        lightmodeButton.pack(pady=5)
        
        # BACK BUTTON
        ttk.Button(self.settings_frame, text="Back", command=self.back_to_main, style='Custom.TButton').pack(side=tk.TOP, padx=10, pady=10)
        # # APPLY BUTTON
        # ttk.Button(self.settings_frame, text="Apply", command=self.confirm_apply, style='Custom.TButton').pack(
        #     side=tk.TOP, padx=10, pady=10)
        
        # Initially hide the settings frame
        self.settings_frame.pack_forget()
    
    def show_settings(self):
        self.main_frame.pack_forget()
        self.settings_frame.pack(fill=tk.BOTH, expand=True)
    
    def show_main(self):
        self.settings_frame.pack_forget()
        self.main_frame.pack(fill=tk.BOTH, expand=True)
    
    def set_theme(self, theme_name):
        self.style.theme_use(theme_name)
        self.current_theme = theme_name
    
    def set_theme_light(self):
        print("Changed to light theme")
        self.set_theme('united')
        self.current_theme = 'united'
    
    def set_theme_dark(self):
        print("Changed to dark theme")
        self.set_theme('darkly')
        self.current_theme = 'darkly'
    
    def back_to_main(self):
        print("Going back to the main menu")
        self.show_main()
    
    def select_files(self):
        file_paths = filedialog.askopenfilenames(
            title='Select files to import',
            filetypes=[('All Files', '*.*')]
        )
        return file_paths
    
    def convert_file_type(self, file_path, new_extension): # Simulate file conversion process (for demo purposes)    
        return f"{file_path}.{new_extension}"
    
    def get_file_name(self, str=""):
        if str == "":
            print("ERROR: Input string empty")
            return
        else:
            slashIndexList = []  # List to hold the indeces of the slashes within the input filepath string
            index = 0  # Index counter for iteration
            for char in str:
                if char == "/":
                    slashIndexList.append(index)
                else:
                    pass
                index += 1
            if slashIndexList != []:  # If the slash list ain't empty
                lastSlash = slashIndexList[-1]
                returnStr = str[lastSlash + 1:]
                return returnStr
    
    def delete_single_row(self, item_id):
        """Delete a row and its associated widgets"""
        print(f"Attempting to delete item_id: {item_id} (type: {type(item_id)})")
        item_id = str(item_id) # Convert item_id to string to ensure consistency
        
        # 1. Remove the delete button if it exists
        if item_id in self.delete_buttons:
            btn = self.delete_buttons[item_id]
            btn.place_forget()
            btn.destroy()
            del self.delete_buttons[item_id]
        
        # 2. Remove the dropdown widget if it exists
        if item_id in self.dropdown_widgets:
            dropdown = self.dropdown_widgets[item_id]
            dropdown.place_forget()
            dropdown.destroy()
            del self.dropdown_widgets[item_id]
        
        # 3. Remove the download button if it exists
        if item_id in self.download_buttons:
            btn = self.download_buttons[item_id]
            btn.place_forget()
            btn.destroy()
            del self.download_buttons[item_id]
        
        # 4. Delete the row from Treeview
        self.file_list.delete(item_id)
    
    def _place_widgets(self, item_id):
        """Place the widgets after the row is rendered."""
        try:
            bbox = self.file_list.bbox(item_id)
            if bbox is None:  # Item not yet rendered
                self.root.after(50, lambda: self._place_widgets(item_id))
                return
                
            x, y, width, height = bbox
            
            # Get column positions
            col_positions = []
            for col in ["#1", "#2", "#3", "#4", "#5"]:
                col_bbox = self.file_list.bbox(item_id, col)
                if col_bbox:
                    col_positions.append(col_bbox[0])
            
            if len(col_positions) >= 5:
                # Place dropdown in "Additional Info" column
                dropdown = self.dropdown_widgets[item_id]
                dropdown.place(x=col_positions[2] + self.file_list.winfo_x(), 
                            y=y + self.file_list.winfo_y() + 2)
                
                # Place download button in "Download" column
                download_btn = self.download_buttons[item_id]
                download_btn.place(x=col_positions[3] + self.file_list.winfo_x(), 
                                y=y + self.file_list.winfo_y() + 2)
                
                # Place delete button in "Delete" column
                delete_btn = self.delete_buttons[item_id]
                delete_btn.place(x=col_positions[4] + self.file_list.winfo_x(), 
                                y=y + self.file_list.winfo_y() + 2)
        
        except (tk.TclError, KeyError) as e:
            print(f"Failed to place widgets for {item_id}: {e}")
    
    def on_file_import(self):  # THIS FUNCTION DOES NOT CONVERT ANYTHING BUT RATHER HANDLES FILE SELECTION!!!
        # Clear existing entries AND their delete buttons in the Treeview
        self.clear_all_entries()
        
        print("=============================================")
        print(f"on_file_import has been called")
        files = self.select_files() # Obtain the currently selected files from the user
        
        # Error check for import of unsupported filetypes
        for file in files: # Iterate through all added files
            currentFileType = os.path.splitext(file)[1]  # The current file's extension
            if currentFileType.lower() not in self.valid_extensions: # If the current file's extension is not valid
                print(f"ERROR: Selected file: {os.path.basename(file)} has unsupported filetype!")
                messagebox.showwarning("Error", f"Selected file: {os.path.basename(file)} has unsupported filetype!")
                return
        
        for file in files:  # Save all uploaded files to the internal list
            self.selected_file_list.append(file)
        
        # String cut the filepath to only include the filename and the file extensions
        
        if ( len(files) > 10 ): # Error check for too many files at once
            messagebox.showwarning("Error", "Please select no more than 10 files for upload.")
            return
        
        if not files:
            messagebox.showwarning("Error", "Please select at least one file.")
            return
        
        conversion_type = self.conversion_type_var.get()
        if not conversion_type:
            messagebox.showwarning("Error", "Please select a conversion type.")
            return
        
        results = {file: self.convert_file_type(file, conversion_type) for file in files}
        
        # Add new entries to the Treeview
        for idx, (file, new_file) in enumerate(results.items(), start=1):  # VERY IMPORTANT: THE LOOP THAT POPULATES THE TABLE IN THE UI
            item_id = self.file_list.insert("", "end", values=(idx, self.get_file_name(file), "Select Filetype", "", ""))
            
            # Create delete button
            delete_btn = ttk.Button(
                self.file_list_frame,
                text="❌",
                command=lambda id=item_id: self.delete_single_row(id),
                style='danger.TButton',
                width=2
            )
            
            # Create dropdown for file type selection
            dropdown_var = tk.StringVar(value="Select Filetype")
            dropdown = ttk.Combobox(
                self.file_list_frame,
                textvariable=dropdown_var,
                values=self.valid_extensions,
                state='readonly',
                width=12
            )
            # Binding that triggers when the user selects a new dropdown menu value
            dropdown.bind('<<ComboboxSelected>>', lambda event, id=item_id: self.on_dropdown_change(event, id))
            
            # Create download button (initially disabled)
            download_btn = ttk.Button(
                self.file_list_frame,
                text="Download",
                command=lambda f=file, id=item_id: self.download_single_file(f, id),
                style='info.TButton',
                width=8,
                state='disabled'
            )
            
            # Store widget references
            self.delete_buttons[item_id] = delete_btn
            self.dropdown_widgets[item_id] = dropdown
            self.download_buttons[item_id] = download_btn
            
            # Place widgets immediately after Treeview updates
            self.file_list.update_idletasks()
            self._place_widgets(item_id)
        
        print(f"selected_file_list: {self.selected_file_list}")
        print(f"file_list: ", self.file_list.get_children())
        print("=============================================")
    
    def clear_all_entries(self):
        print("clear_all_entries was called")
        """
        Clears all uploaded files AND their widgets
        """
        # Clear all widgets first
        for button in self.delete_buttons.values():
            button.destroy()
        for dropdown in self.dropdown_widgets.values():
            dropdown.destroy()
        for button in self.download_buttons.values():
            button.destroy()
            
        self.delete_buttons.clear()
        self.dropdown_widgets.clear()
        self.download_buttons.clear()
        
        # Then clear the treeview
        for item in self.file_list.get_children():
            self.file_list.delete(item)
        
        # Delete the internal selected file list
        self.selected_file_list = []
        self.converted_files.clear()
    
    def on_configure(self, event): # DO NOT REMOVE 'event' FROM THE ARGS, IT CRASHES ENTIRE PROGRAM!!!
        self.root.update_idletasks()
        self.root_x = self.root.winfo_rootx()
        self.root_y = self.root.winfo_rooty()
        # print(f"Window moved. New root position: ({self.root_x}, {self.root_y})")
    
    def on_dropdown_change(self, event, item_id):
        """Handle dropdown selection change - reset conversion state"""
        dropdown = event.widget
        selected_value = dropdown.get()
        print(f"Dropdown for item {item_id} changed to: {selected_value}")
        
        # Reset the conversion state for this item
        self.reset_item_conversion_state(item_id)
    
    def reset_item_conversion_state(self, item_id):
        """Reset conversion state for a specific item"""
        # Disable the download button
        if item_id in self.download_buttons:
            self.download_buttons[item_id].configure(state='disabled')
        
        # Find and remove the converted file from converted_files dict
        # We need to find which original file corresponds to this item_id
        values = self.file_list.item(item_id, 'values')
        if values:
            file_index = int(values[0]) - 1  # Convert to 0-based index
            if 0 <= file_index < len(self.selected_file_list):
                original_file = self.selected_file_list[file_index]
                if original_file in self.converted_files:
                    # Delete the temp converted file to not waste space
                    converted_file_path = self.converted_files[original_file]
                    if os.path.exists(converted_file_path):
                        try:
                            os.remove(converted_file_path)
                            print(f"Deleted temp file: {converted_file_path}")
                        except Exception as e:
                            print(f"Error deleting temp file: {e}")
                    
                    # Remove from converted_files dict
                    del self.converted_files[original_file]
                    print(f"Reset conversion state for: {original_file}")
    
    def text_file_to_pdf(self, text_file_path, output_pdf_path):
        # Create a canvas object
        c = canvas.Canvas(output_pdf_path, pagesize=letter)
        width, height = letter
        
        # Set the font and size
        font_name = "Helvetica"
        font_size = 12
        c.setFont(font_name, font_size)
        
        # Define the margin and line height
        margin = 72  # 1 inch margin
        line_height = font_size * 1.2  # line height (can be adjusted)
        
        # Read the text file
        with open(text_file_path, 'r') as file:
            text = file.readlines()
        
        # Calculate the maximum number of lines per page
        max_lines_per_page = int((height - 2 * margin) / line_height)
        
        # Draw the text on the PDF
        y = height - margin
        for i, line in enumerate(text):
            if i % max_lines_per_page == 0 and i != 0:
                c.showPage()  # start a new page
                c.setFont(font_name, font_size)
                y = height - margin
            
            c.drawString(margin, y, line.strip())
            y -= line_height
        
        # Save the PDF
        c.save()
    
    def docx_to_pdf(self, docx_file_path, output_pdf_path):
        # Convert the DOCX file to PDF
        convert(docx_file_path, output_pdf_path)
    
    def convert_all_files(self):
        """Convert all files based on their individual dropdown selections"""
        print("==================================================")
        print(f"convert_all_files function has been called")
        print(f"selected_file_list: {self.selected_file_list}")
        print(f"file_list: {self.file_list.get_children()}")
        print("==================================================")
        if not self.selected_file_list:
            messagebox.showwarning("Error", "No files to convert.")
            return
        
        converted_count = 0
        
        # Create a mapping of item_id to original file path
        item_to_file_mapping = {}
        for idx, item_id in enumerate(self.file_list.get_children()):
            if idx < len(self.selected_file_list):
                item_to_file_mapping[item_id] = self.selected_file_list[idx]
        
        for item_id in self.file_list.get_children():
            if item_id in self.dropdown_widgets and item_id in item_to_file_mapping:
                dropdown = self.dropdown_widgets[item_id]
                selected_extension = dropdown.get()
                
                if selected_extension == "Select Filetype" or not selected_extension:
                    print(f"Skipping item {item_id}: no filetype selected")
                    continue
                
                print(f"Processing item {item_id} with extension {selected_extension}")
                
                original_file = item_to_file_mapping[item_id]
                print(f"Converting file: {original_file}")
                converted_file_path = self.convert_single_file(original_file, selected_extension)
                
                if converted_file_path:
                    self.converted_files[original_file] = converted_file_path
                    # Enable the download button
                    if item_id in self.download_buttons:
                        self.download_buttons[item_id].configure(state='normal')
                    converted_count += 1
                    print(f"Successfully converted: {original_file}")
                else:
                    print(f"Failed to convert: {original_file}")
        
        if converted_count > 0:
            messagebox.showinfo("Success", f"Successfully converted {converted_count} file(s).")
        else:
            messagebox.showwarning("Warning", "No files were converted. Please select output filetypes.")
    
    def convert_single_file(self, file_path, conversion_extension):
        """Convert a single file and return the output path"""
        try:
            currentFileType = os.path.splitext(file_path)[1]
            base_name = os.path.splitext(os.path.basename(file_path))[0]
            output_file = os.path.join(self.temp_dir, base_name + conversion_extension)
            
            # Image file conversions
            img_file_extensions = ['.bmp', '.jpg', '.jpeg', '.png', '.heic', '.pdf']
            text_file_extensions = ['.txt', '.docx']
            
            if currentFileType.lower() in [ext.lower() for ext in img_file_extensions]:
                if conversion_extension.lower() not in [ext.lower() for ext in text_file_extensions]:
                    if currentFileType.lower() == conversion_extension.lower():
                        return None  # Same format, no conversion needed
                    
                    img = Image.open(file_path)
                    img.save(output_file, conversion_extension.replace(".", "").upper())
                    return output_file
            
            # Text file conversions
            elif currentFileType.lower() in [ext.lower() for ext in text_file_extensions]:
                if currentFileType.lower() == conversion_extension.lower():
                    return None  # Same format, no conversion needed
                
                if conversion_extension.lower() == '.pdf':
                    if currentFileType.lower() == '.txt':
                        self.text_file_to_pdf(file_path, output_file)
                        return output_file
                    elif currentFileType.lower() == '.docx':
                        self.docx_to_pdf(file_path, output_file)
                        return output_file
                
                elif conversion_extension.lower() == '.txt' and currentFileType.lower() == '.docx':
                    pypandoc.convert_file(file_path, 'plain', outputfile=output_file)
                    return output_file
                
                elif conversion_extension.lower() == '.docx' and currentFileType.lower() == '.txt':
                    document = Document()
                    with open(file_path, 'r', encoding='utf-8') as myfile:
                        myfile_content = myfile.read()
                    myfile_content = re.sub(r'[^\x00-\x7F]+|\x0c', ' ', myfile_content)
                    document.add_heading(os.path.basename(file_path), 0)
                    document.add_paragraph(myfile_content)
                    document.save(output_file)
                    return output_file
            
            return None
            
        except Exception as e:
            print(f"Error converting {file_path}: {e}")
            messagebox.showerror("Error", f"Failed to convert {os.path.basename(file_path)}: {str(e)}")
            return None
    
    def download_single_file(self, original_file, item_id):
        """Download a single converted file"""
        if original_file in self.converted_files:
            converted_file = self.converted_files[original_file]
            if os.path.exists(converted_file):
                save_path = filedialog.asksaveasfilename(
                    defaultextension=os.path.splitext(converted_file)[1],
                    filetypes=[('All Files', '*.*')],
                    initialfile=os.path.basename(converted_file)
                )
                if save_path:
                    try:
                        shutil.copy2(converted_file, save_path)
                        messagebox.showinfo("Success", f"File saved to {save_path}")
                    except Exception as e:
                        messagebox.showerror("Error", f"Failed to save file: {str(e)}")
            else:
                messagebox.showerror("Error", "Converted file not found. Please convert the file first.")
        else:
            messagebox.showwarning("Warning", "File not converted yet. Please convert the file first.")
    
    def download_all_files(self):
        """Download all converted files to a selected directory"""
        if not self.converted_files:
            messagebox.showwarning("Warning", "No converted files available for download.")
            return
        
        save_dir = filedialog.askdirectory(title="Select directory to save all converted files")
        if save_dir:
            try:
                saved_count = 0
                for original_file, converted_file in self.converted_files.items():
                    if os.path.exists(converted_file):
                        dest_path = os.path.join(save_dir, os.path.basename(converted_file))
                        shutil.copy2(converted_file, dest_path)
                        saved_count += 1
                
                messagebox.showinfo("Success", f"Successfully saved {saved_count} file(s) to {save_dir}")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to save files: {str(e)}")
    
    def cleanup_temp_files(self):
        """Clean up temporary directory and files"""
        try:
            if hasattr(self, 'temp_dir') and os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
        except Exception as e:
            print(f"Error cleaning up temp files: {e}")
    
    def __del__(self):
        """Destructor to clean up temp files"""
        self.cleanup_temp_files()


if __name__ == "__main__":
    root = ttk.Window(title="Pyfile Converter", themename="united")  # Set initial theme
    root.geometry("740x300")  # Set window size
    app = FileConverterApp(root)
    root.mainloop()
